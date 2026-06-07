"""
CV parser — extracts plain text from PDF, DOCX, or raw text uploads.

Requires:
  pip install pymupdf python-docx

Both are optional: if not installed, falls back to raw text.
"""

import io
import logging
import re

logger = logging.getLogger(__name__)


# ── PDF ─────���─────────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file (bytes). Returns plain text."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages)
    except ImportError:
        logger.warning(
            "PyMuPDF not installed. Install with: pip install pymupdf\n"
            "Falling back to empty string for PDF."
        )
        return ""
    except Exception as exc:
        logger.error("PDF parsing failed: %s", exc)
        return ""


# ── DOCX ──────────────────────────────────────────────────────────────────────

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file (bytes). Returns plain text."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except ImportError:
        logger.warning(
            "python-docx not installed. Install with: pip install python-docx\n"
            "Falling back to empty string for DOCX."
        )
        return ""
    except Exception as exc:
        logger.error("DOCX parsing failed: %s", exc)
        return ""


# ── Dispatcher ───────────��────────────────────────────────────────────────────

def parse_uploaded_file(file_obj) -> str:
    """
    Accept a Streamlit UploadedFile object and return extracted plain text.
    Handles .pdf, .docx, .doc (text extraction only), and .txt.
    """
    if file_obj is None:
        return ""

    name = getattr(file_obj, "name", "").lower()
    raw_bytes = file_obj.read()

    if name.endswith(".pdf"):
        text = parse_pdf(raw_bytes)
    elif name.endswith(".docx"):
        text = parse_docx(raw_bytes)
    elif name.endswith(".txt"):
        try:
            text = raw_bytes.decode("utf-8", errors="replace")
        except Exception:
            text = ""
    else:
        # Try as text anyway
        try:
            text = raw_bytes.decode("utf-8", errors="replace")
        except Exception:
            text = ""

    return clean_cv_text(text)


# ── Text cleaning ───────���─────────────────────────────���───────────────────────

def clean_cv_text(text: str) -> str:
    """Normalise whitespace and remove junk characters from extracted CV text."""
    if not text:
        return ""
    # Remove control characters except newline/tab
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\x80-\xFF]", " ", text)
    # Collapse runs of spaces (not newlines)
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse more than 3 consecutive blank lines to 2
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


# ── Structured extraction ─────────────────────────────────────────────────────

def extract_cv_sections(cv_text: str) -> dict[str, str]:
    """
    Best-effort section extraction: summary, experience, education, skills.
    Returns a dict; values may be empty if section not found.
    """
    sections: dict[str, str] = {
        "summary":    "",
        "experience": "",
        "education":  "",
        "skills":     "",
        "full":       cv_text,
    }
    if not cv_text:
        return sections

    # Common section header patterns (case-insensitive)
    _HEADERS = {
        "summary":    r"(profile|summary|objective|about me|career summary)",
        "experience": r"(experience|employment|work history|career history|positions? held)",
        "education":  r"(education|qualifications?|academic|degrees?|study)",
        "skills":     r"(skills|technical|competenc|expertise|capabilities)",
    }

    lines = cv_text.split("\n")
    current_section = None
    buffer: list[str] = []

    for line in lines:
        stripped = line.strip()
        matched = False
        for section, pattern in _HEADERS.items():
            if re.match(pattern, stripped, re.IGNORECASE) and len(stripped) < 60:
                if current_section and buffer:
                    sections[current_section] = "\n".join(buffer).strip()
                current_section = section
                buffer = []
                matched = True
                break
        if not matched and current_section:
            buffer.append(line)

    if current_section and buffer:
        sections[current_section] = "\n".join(buffer).strip()

    return sections


def extract_years_experience(cv_text: str) -> int:
    """
    Estimate years of experience from the CV text.
    Looks for patterns like '5 years', '10+ years', etc.
    Returns 0 if nothing found.
    """
    patterns = [
        r"(\d+)\+?\s*years?\s+(?:of\s+)?(?:professional\s+)?experience",
        r"experience\s+(?:of\s+)?(\d+)\+?\s*years?",
        r"(\d+)\s*(?:-|–)\s*\d+\s*years?",
    ]
    for pat in patterns:
        m = re.search(pat, cv_text, re.IGNORECASE)
        if m:
            try:
                return int(m.group(1))
            except ValueError:
                pass
    return 0
